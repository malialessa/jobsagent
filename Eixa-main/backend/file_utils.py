import base64
import re
import io
import logging
from PIL import Image
import fitz # Importação para PDF (PyMuPDF)
from docx import Document # Importação para DOCX (python-docx)
from typing import Dict

logger = logging.getLogger(__name__)

# --- Constantes de Configuração ---
MAX_FILE_SIZE_BYTES = 10 * 1024 * 1024 # 10 MB
MAX_PDF_PAGES = 50 # Limite para evitar abuso de processamento

def process_uploaded_file(base64_data: str, filename: str, mimetype: str) -> Dict:
    """
    Processa um arquivo enviado (base64), sanitiza a entrada e extrai o conteúdo relevante.
    Retorna um dicionário estruturado com tipo, conteúdo e metadados.

    Known Limitations:
    - PDF Processing: Does not perform Optical Character Recognition (OCR). PDFs containing
      only scanned images will not yield extractable text.
    - DOCX Processing: Does not extract images or other embedded objects from DOCX files,
      only textual content from paragraphs and tables.
    - File Types: Does not support specialized formats like RAW, SVG, or older .doc files.
    """
    if not base64_data:
        raise ValueError("Dados em base64 não podem estar vazios.")

    # 1. Sanitização: Remove o prefixo 'data:...' comum em uploads de frontend
    try:
        base64_clean = re.sub(r'^data:.+;base64,', '', base64_data)
        decoded_bytes = base64.b64decode(base64_clean)
        file_size = len(decoded_bytes)
    except (TypeError, ValueError) as e:
        logger.error(f"Falha ao decodificar base64 para o arquivo '{filename}': {e}", exc_info=True)
        raise ValueError("Dados em base64 inválidos ou corrompidos.")

    if file_size > MAX_FILE_SIZE_BYTES:
        raise ValueError(f"O tamanho do arquivo excede o limite de {MAX_FILE_SIZE_BYTES / (1024*1024):.1f} MB.")

    metadata = {
        "filename": filename,
        "size_bytes": file_size,
        "mime_type": mimetype
    }

    # 2. Processamento por Tipo de Arquivo
    if mimetype.startswith('image/'):
        try:
            # Apenas verifica se é uma imagem válida sem processar o conteúdo
            Image.open(io.BytesIO(decoded_bytes))
            logger.info(f"Arquivo de imagem processado: '{filename}' ({mimetype})")
            return {
                'type': 'image',
                'content': {'base64_image': base64_clean, 'mime_type': mimetype},
                'metadata': metadata
            }
        except Exception as e:
            logger.error(f"Arquivo de imagem inválido '{filename}': {e}", exc_info=True)
            raise ValueError(f"Não foi possível processar o arquivo de imagem: {filename}")

    elif mimetype == 'application/pdf':
        try:
            doc = fitz.open(stream=decoded_bytes, filetype="pdf")
            text_content = ""
            for i, page in enumerate(doc):
                if i >= MAX_PDF_PAGES:
                    logger.warning(f"PDF '{filename}' excedeu o limite de {MAX_PDF_PAGES} páginas. Processamento interrompido.")
                    break
                text_content += page.get_text("text")
            doc.close()

            if not text_content.strip():
                logger.warning(f"PDF '{filename}' não contém texto extraível. Pode ser um arquivo de imagem escaneado ou sem texto selecionável.")
                return {
                    'type': 'text',
                    'content': {'text_content': "[AVISO: O PDF não contém texto legível e pode ser uma imagem escaneada ou sem texto selecionável.]"},
                    'metadata': metadata
                }

            logger.info(f"PDF processado: '{filename}', extraídos {len(text_content)} caracteres.")
            return {
                'type': 'text',
                'content': {'text_content': text_content},
                'metadata': metadata
            }
        except Exception as e:
            logger.error(f"Erro ao processar PDF '{filename}' com PyMuPDF: {e}", exc_info=True)
            raise ValueError(f"Não foi possível processar o arquivo PDF: {filename}")

    elif mimetype == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
        try:
            doc = Document(io.BytesIO(decoded_bytes))
            text_content_parts = []
            for para in doc.paragraphs:
                text_content_parts.append(para.text)
            # Adiciona texto de tabelas, se houver
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content_parts.append(cell.text)

            text_content = "\n".join(text_content_parts)
            
            if not text_content.strip():
                logger.warning(f"DOCX '{filename}' não contém texto extraível. Pode ser um arquivo vazio ou com conteúdo não textual.")
                return {
                    'type': 'text',
                    'content': {'text_content': "[AVISO: O DOCX não contém texto legível ou pode estar vazio.]"},
                    'metadata': metadata
                }

            logger.info(f"Arquivo DOCX processado: '{filename}', extraídos {len(text_content)} caracteres.")
            return {
                'type': 'text',
                'content': {'text_content': text_content},
                'metadata': metadata
            }
        except Exception as e:
            logger.error(f"Erro ao processar DOCX '{filename}': {e}", exc_info=True)
            raise ValueError(f"Não foi possível processar o arquivo DOCX: {filename}")

    else:
        logger.warning(f"Tipo de arquivo não suportado: '{mimetype}' para o arquivo '{filename}'.")
        raise ValueError(f"Tipo de arquivo não suportado: {mimetype}. Apenas imagens (JPG/PNG), PDF e DOCX são permitidos.")